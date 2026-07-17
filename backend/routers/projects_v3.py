"""
V3 project/task/development-point API for agent-driven development iteration.
"""

import os
import asyncio
import json
import re
import shlex
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from fastapi import APIRouter, HTTPException, Query
from pydantic import BaseModel, Field

from data_manager import data_manager
from project_manager import project_manager
from knowledge_manager import knowledge_manager
from openclaw_integration import openclaw_integration
from unified_data_manager import unified_data_manager
from services.project_task_sync import sync_v3_task_to_v2
from services.work_run_service import (
    InvalidWorkTransition,
    WorkRunLeaseConflict,
    WorkRunNotFound,
    work_run_service,
)
import skill_manager


router = APIRouter(prefix="/api/v3", tags=["projects-v3"])

BASE_DIR = Path(__file__).resolve().parents[1]
AGENT_ORGANIZATION_FILE = BASE_DIR / "data" / "agent_organization.json"
AGENT_TEMPLATES_FILE = BASE_DIR / "data" / "agent_templates.json"
AGENT_INSTANCES_FILE = BASE_DIR / "data" / "agent_instances.json"
AGENT_RUNNER_FILE = BASE_DIR / "data" / "agent_runner.json"
OPENCLAW_WORKSPACE = Path.home() / ".openclaw" / "workspace"
OPENCLAW_AGENTS_DIR = OPENCLAW_WORKSPACE / "agents"
NINJA_TURTLES_QUEUE_FILE = OPENCLAW_AGENTS_DIR / "ninja-turtles" / "dev-loop" / "queue.json"
DEFAULT_CODEX_BIN = os.getenv("CODEX_BIN", "/opt/homebrew/bin/codex")
DEFAULT_OPENCLAW_BIN = os.getenv("OPENCLAW_BIN", "/opt/homebrew/opt/node/bin/openclaw")
DEFAULT_CODEX_PATH = os.getenv(
    "CODEX_PATH",
    "/opt/homebrew/bin:/opt/homebrew/opt/node/bin:/opt/homebrew/Cellar/node/25.6.1_1/bin:/usr/local/bin:/usr/bin:/bin:/usr/sbin:/sbin",
)
DEFAULT_TEAM_DASHBOARD_REPO = OPENCLAW_WORKSPACE / "team-dashboard"
DEFAULT_CODEX_REPO = os.getenv(
    "CODEX_DEFAULT_REPO",
    str(DEFAULT_TEAM_DASHBOARD_REPO if DEFAULT_TEAM_DASHBOARD_REPO.exists() else BASE_DIR.parent),
)
_ELASTIC_RUNNER_TASK: Optional[asyncio.Task] = None
_ELASTIC_RUNNER_STOP: Optional[asyncio.Event] = None
_ELASTIC_RUNNER_STATE: dict[str, Any] = {
    "enabled": False,
    "running": False,
    "last_tick_at": None,
    "last_claimed": False,
    "last_error": "",
}


class DesignDocument(BaseModel):
    id: Optional[str] = None
    project_id: Optional[str] = None
    version: int = 1
    status: str = "draft"
    summary: str = ""
    usage_requirements: list[Any] = Field(default_factory=list)
    data_structure: dict[str, Any] = Field(default_factory=dict)
    system_architecture: dict[str, Any] = Field(default_factory=dict)
    system_functions: list[Any] = Field(default_factory=list)
    api_interfaces: list[Any] = Field(default_factory=list)
    task_breakdown_guidance: list[Any] = Field(default_factory=list)
    parallel_tasks: list[Any] = Field(default_factory=list)
    risks: list[Any] = Field(default_factory=list)
    changelog: list[Any] = Field(default_factory=list)
    author_agent: str = "project-manager"
    approved_by: str = ""
    approved_at: Optional[str] = None
    change_summary: Optional[str] = None


class DesignDocumentUpdate(BaseModel):
    status: Optional[str] = None
    summary: Optional[str] = None
    usage_requirements: Optional[list[Any]] = None
    data_structure: Optional[dict[str, Any]] = None
    system_architecture: Optional[dict[str, Any]] = None
    system_functions: Optional[list[Any]] = None
    api_interfaces: Optional[list[Any]] = None
    task_breakdown_guidance: Optional[list[Any]] = None
    parallel_tasks: Optional[list[Any]] = None
    risks: Optional[list[Any]] = None
    author_agent: Optional[str] = None
    change_summary: Optional[str] = None


class DesignDocumentRevise(BaseModel):
    agent_id: str = "project-manager"
    change_summary: str = "项目设计文档修订"
    status: str = "draft"
    updates: dict[str, Any] = Field(default_factory=dict)


class ApproveDesignDocumentRequest(BaseModel):
    agent_id: str = "project-manager"


class ProjectCreate(BaseModel):
    id: Optional[str] = None
    name: str
    project_type: str = "software"
    type: Optional[str] = None
    description: str = ""
    status: str = "planning"
    priority: str = "medium"
    owner_agent: str = "optimus"
    progress: float = 0
    current_phase: str = "planning"
    context: dict[str, Any] = Field(default_factory=dict)
    enabled_modules: list[str] = Field(default_factory=list)
    product_bindings: list[dict[str, Any]] = Field(default_factory=list)
    design_doc: Optional[DesignDocument] = None
    document_spec: dict[str, Any] = Field(default_factory=dict)


class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    project_type: Optional[str] = None
    type: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    priority: Optional[str] = None
    owner_agent: Optional[str] = None
    progress: Optional[float] = None
    current_phase: Optional[str] = None
    context: Optional[dict[str, Any]] = None
    enabled_modules: Optional[list[str]] = None
    product_bindings: Optional[list[dict[str, Any]]] = None
    document_spec: Optional[dict[str, Any]] = None


class DevelopmentPointCreate(BaseModel):
    id: Optional[str] = None
    task_id: Optional[str] = None
    title: str
    description: str = ""
    status: str = "todo"
    weight: float = 1
    completion_evidence: str = ""
    checklist: list[str] = Field(default_factory=list)
    assigned_agent: str = ""


class DevelopmentPointUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    weight: Optional[float] = None
    completion_evidence: Optional[str] = None
    checklist: Optional[list[str]] = None
    assigned_agent: Optional[str] = None


class TaskCreate(BaseModel):
    id: Optional[str] = None
    type: str = "development"
    title: str
    description: str = ""
    assignee_agent: str = ""
    assignee_agent_id: str = ""
    status: str = "todo"
    progress: float = 0
    priority: str = "medium"
    dependencies: list[str] = Field(default_factory=list)
    acceptance_criteria: list[str] = Field(default_factory=list)
    context: dict[str, Any] = Field(default_factory=dict)
    result_summary: str = ""
    development_points: list[DevelopmentPointCreate] = Field(default_factory=list)


class TaskUpdate(BaseModel):
    type: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    assignee_agent: Optional[str] = None
    assignee_agent_id: Optional[str] = None
    status: Optional[str] = None
    progress: Optional[float] = None
    priority: Optional[str] = None
    dependencies: Optional[list[str]] = None
    acceptance_criteria: Optional[list[str]] = None
    context: Optional[dict[str, Any]] = None
    result_summary: Optional[str] = None



class AgentSkillsUpdate(BaseModel):
    skill_ids: list[str] = Field(default_factory=list)


class AssignRequest(BaseModel):
    assignee_agent: str


class LogCreate(BaseModel):
    task_id: Optional[str] = None
    agent_id: str = "system"
    action: str = "note"
    content: str


class ProjectChatRequest(BaseModel):
    agent_id: str = "optimus"
    message: str
    intent: str = "chat"
    task_id: str = ""
    chapter_id: str = ""
    role: str = "user"
    attachments: list[dict[str, Any]] = Field(default_factory=list)


class ProjectAgentActionRequest(BaseModel):
    agent_id: str = "optimus"
    action_type: str = "note"
    instruction: str = ""
    task_id: str = ""
    chapter_id: str = ""
    task_title: str = ""
    task_type: str = ""
    assignee_agent: str = ""
    payload: dict[str, Any] = Field(default_factory=dict)


class SoftwareSpecUpdate(BaseModel):
    requirements: Optional[list[Any]] = None
    design_doc: Optional[dict[str, Any]] = None
    architecture: Optional[dict[str, Any]] = None
    database_design: Optional[dict[str, Any]] = None
    api_design: Optional[list[Any]] = None
    frontend_design: Optional[dict[str, Any]] = None
    test_plan: Optional[list[Any]] = None
    deployment_plan: Optional[list[Any]] = None
    agent_id: str = "project-manager"


class DocumentSectionCreate(BaseModel):
    id: Optional[str] = None
    parent_id: str = ""
    title: str
    summary: str = ""
    main_content: str = ""
    content_brief: str = ""
    key_points: list[str] = Field(default_factory=list)
    images: list[Any] = Field(default_factory=list)
    status: str = "planning"
    assigned_agent: str = ""
    assigned_agent_id: str = ""
    order_index: Optional[int] = None
    agent_id: str = "project-manager"


class DocumentSectionUpdate(BaseModel):
    parent_id: Optional[str] = None
    title: Optional[str] = None
    summary: Optional[str] = None
    main_content: Optional[str] = None
    content_brief: Optional[str] = None
    key_points: Optional[list[str]] = None
    images: Optional[list[Any]] = None
    status: Optional[str] = None
    assigned_agent: Optional[str] = None
    assigned_agent_id: Optional[str] = None
    order_index: Optional[int] = None
    agent_id: str = "project-manager"


class DocumentAssetCreate(BaseModel):
    id: Optional[str] = None
    chapter_id: str = ""
    section_id: str = ""
    chapter_title: str = ""
    type: str = "image"
    title: str
    description: str = ""
    file_path: str = ""
    status: str = "planned"
    order_index: Optional[int] = None
    agent_id: str = "project-manager"


class DocumentAssetUpdate(BaseModel):
    chapter_id: Optional[str] = None
    section_id: Optional[str] = None
    chapter_title: Optional[str] = None
    type: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    file_path: Optional[str] = None
    status: Optional[str] = None
    order_index: Optional[int] = None
    agent_id: str = "project-manager"


class DocumentWorkdraftSyncRequest(BaseModel):
    source_word_path: str = ""
    force: bool = False
    agent_id: str = "project-manager"


class DecomposeRequest(BaseModel):
    reasoning_summary: str = ""
    agent_id: str = "project-manager"
    new_tasks: list[TaskCreate] = Field(default_factory=list)
    updated_tasks: list[dict[str, Any]] = Field(default_factory=list)
    project_updates: dict[str, Any] = Field(default_factory=dict)
    new_development_points: list[DevelopmentPointCreate] = Field(default_factory=list)
    updated_development_points: list[dict[str, Any]] = Field(default_factory=list)


class CompletePointRequest(BaseModel):
    agent_id: str
    completion_evidence: str = ""
    result_summary: str = ""


class PointTransitionRequest(BaseModel):
    agent_id: str = ""
    reason: str = ""
    completion_evidence: str = ""
    result_summary: str = ""


class KnowledgeLinkRequest(BaseModel):
    node_id: str
    title: str = ""
    type: str = ""
    path: str = ""
    relation: str = "related"
    reason: str = ""
    confirmed_by: str = "project-manager"


class AgentSpawnRequest(BaseModel):
    requested_by: str
    base_agent_id: str
    project_id: str = ""
    task_id: str = ""
    development_point_id: str = ""
    reason: str = ""
    release_policy: str = "task_completed"
    instance_id: Optional[str] = None


class AgentReleaseRequest(BaseModel):
    released_by: str = "scheduler"
    reason: str = "manual_release"


class AgentCompleteRequest(BaseModel):
    completed_by: str = "scheduler"
    result_summary: str = ""
    memory_summary: str = ""
    decisions: list[str] = Field(default_factory=list)
    blockers: list[str] = Field(default_factory=list)
    next_actions: list[str] = Field(default_factory=list)


class AgentRunnerTickRequest(BaseModel):
    instance_id: Optional[str] = None
    execute: bool = False
    max_seconds: int = Field(default=600, ge=10, le=3600)


class WorkRunTransitionRequest(BaseModel):
    status: str
    actor: str = "human-reviewer"
    detail: str = ""
    result_summary: Optional[str] = None


def _model_dict(model: BaseModel, exclude_unset: bool = False) -> dict:
    return model.model_dump(exclude_unset=exclude_unset)


def _safe_id(value: str) -> str:
    return "".join(ch if ch.isalnum() or ch in {"-", "_"} else "-" for ch in str(value).strip().lower()).strip("-")


def _safe_filename(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in {"-", "_", " ", "."} else "-" for ch in str(value).strip())
    return cleaned.strip(" .-") or "document-project"


def _knowledge_root() -> Path:
    return knowledge_manager.vault_path.resolve()


def _resolve_vault_file(raw_path: str) -> Path:
    if not raw_path:
        raise HTTPException(status_code=400, detail="source_word_path is required")
    root = _knowledge_root()
    path = Path(os.path.expanduser(raw_path))
    if not path.is_absolute():
        path = root / path
    try:
        resolved = path.resolve()
    except OSError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid source_word_path: {raw_path}") from exc
    if resolved != root and root not in resolved.parents:
        raise HTTPException(status_code=400, detail="source_word_path must be inside the knowledge vault")
    if not resolved.exists() or not resolved.is_file():
        raise HTTPException(status_code=404, detail=f"Word source not found: {raw_path}")
    if resolved.suffix.lower() != ".docx":
        raise HTTPException(status_code=400, detail="当前自动同步仅支持 .docx，旧 .doc 请先另存为 .docx")
    return resolved


def _relative_to_vault(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(_knowledge_root()))
    except ValueError:
        return str(path)


def _score_word_candidate(path: Path, project: dict[str, Any]) -> int:
    text = str(path).lower()
    name = str(project.get("name") or "").lower()
    score = 0
    for token in [name, "博士", "论文", "论文章节", "毕业论文"]:
        if token and token.lower() in text:
            score += 20
    if "10-成果库" in str(path):
        score += 8
    if "已排版" in text:
        score += 8
    if "初稿" in text:
        score += 5
    if "_pandoc_temp" in text or "/tmp/" in text:
        score -= 12
    if path.name.startswith("~$"):
        score -= 100
    try:
        score += min(int(path.stat().st_mtime // 86400), 10_000)
    except OSError:
        pass
    return score


def _find_word_source(project: dict[str, Any], requested_path: str = "") -> Path:
    if requested_path:
        return _resolve_vault_file(requested_path)
    spec = project.get("document_spec") if isinstance(project.get("document_spec"), dict) else {}
    saved_path = (spec.get("source_word") or {}).get("path") if isinstance(spec.get("source_word"), dict) else ""
    if saved_path:
        try:
            return _resolve_vault_file(saved_path)
        except HTTPException:
            pass

    root = _knowledge_root()
    if not root.exists():
        raise HTTPException(status_code=404, detail="知识库目录不存在，无法查找 Word 原文")
    candidates = [p for p in root.rglob("*.docx") if p.is_file() and not p.name.startswith("~$")]
    candidates.sort(key=lambda p: _score_word_candidate(p, project), reverse=True)
    best = candidates[0] if candidates else None
    if not best or _score_word_candidate(best, project) < 20:
        raise HTTPException(status_code=404, detail="未在知识库中找到匹配该文档项目的 .docx 原文，请传入 source_word_path")
    return best


def _heading_level(style_name: str) -> int:
    style = (style_name or "").strip().lower()
    match = re.search(r"(?:heading|标题)\s*([1-6])", style)
    return int(match.group(1)) if match else 0


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    divider = "| " + " | ".join(["---"] * width) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, divider, *body])


def _docx_to_markdown(source_path: Path, project: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(status_code=500, detail="服务器缺少 python-docx，无法解析 Word 原文") from exc

    try:
        document = Document(str(source_path))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Word 原文解析失败: {source_path.name}") from exc

    generated_at = datetime.now(timezone.utc).isoformat()
    lines = [
        "---",
        f"title: {project.get('name') or source_path.stem}",
        f"source_word: {_relative_to_vault(source_path)}",
        f"generated_at: {generated_at}",
        "status: workdraft",
        "---",
        "",
    ]
    paragraph_count = 0
    heading_count = 0
    for para in document.paragraphs:
        text = " ".join(para.text.split())
        if not text:
            continue
        level = _heading_level(getattr(para.style, "name", ""))
        if level:
            lines.append(f"{'#' * level} {text}")
            heading_count += 1
        elif "list" in str(getattr(para.style, "name", "")).lower() or "列表" in str(getattr(para.style, "name", "")):
            lines.append(f"- {text}")
        else:
            lines.append(text)
        lines.append("")
        paragraph_count += 1

    table_count = 0
    for table in document.tables:
        md_table = _table_to_markdown(table)
        if md_table:
            lines.extend([md_table, ""])
            table_count += 1

    markdown = "\n".join(lines).strip() + "\n"
    stats = {
        "paragraph_count": paragraph_count,
        "heading_count": heading_count,
        "table_count": table_count,
        "generated_at": generated_at,
        "size_chars": len(markdown),
    }
    return markdown, stats


def _slug(value: str) -> str:
    value = re.sub(r"[^\w\u4e00-\u9fff]+", "-", value.strip().lower())
    return value.strip("-") or "section"


_CHINESE_NUMBERS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


def _chapter_number(value: str) -> Optional[int]:
    match = re.search(r"第\s*([0-9一二三四五六七八九十]+)\s*章", value or "")
    if not match:
        return None
    raw = match.group(1)
    if raw.isdigit():
        return int(raw)
    if raw == "十":
        return 10
    if raw.startswith("十"):
        return 10 + _CHINESE_NUMBERS.get(raw[-1], 0)
    if raw.endswith("十"):
        return _CHINESE_NUMBERS.get(raw[0], 1) * 10
    return _CHINESE_NUMBERS.get(raw)


def _extract_markdown_headings(markdown: str, sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    project_sections_by_number = {_chapter_number(s.get("title", "")): s for s in sections if _chapter_number(s.get("title", ""))}
    links = []
    for line_number, line in enumerate(markdown.splitlines(), start=1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not match:
            continue
        level = len(match.group(1))
        heading = match.group(2).strip()
        chapter_no = _chapter_number(heading)
        matched_section = project_sections_by_number.get(chapter_no) if chapter_no else None
        if not matched_section:
            normalized_heading = re.sub(r"\s+", "", heading)
            for section in sections:
                normalized_title = re.sub(r"\s+", "", section.get("title", ""))
                if normalized_title and (normalized_title in normalized_heading or normalized_heading in normalized_title):
                    matched_section = section
                    break
        links.append({
            "id": f"md-heading-{line_number}",
            "heading": heading,
            "level": level,
            "anchor": _slug(heading),
            "line": line_number,
            "section_id": matched_section.get("id", "") if matched_section else "",
            "section_title": matched_section.get("title", "") if matched_section else "",
            "source": "markdown_workdraft",
        })
    return links


def _chapters_with_workdraft_outline(project: dict[str, Any], links: list[dict[str, Any]]) -> list[dict[str, Any]]:
    spec = project.get("document_spec") if isinstance(project.get("document_spec"), dict) else {}
    chapters = [dict(item) for item in spec.get("chapters", []) if isinstance(item, dict)]
    if not chapters or not links:
        return chapters
    top_links = [link for link in links if link.get("level") <= 2 and link.get("section_id")]
    for chapter in chapters:
        section_links = [link for link in top_links if link.get("section_id") == chapter.get("id")]
        child_links = [link for link in links if link.get("level", 9) > 1 and link.get("section_id") == chapter.get("id")]
        if child_links:
            chapter["outline_items"] = [
                {
                    "id": f"{chapter.get('id', 'chapter')}-md-{index + 1}",
                    "title": item["heading"],
                    "summary": f"来自 Markdown 工作稿第 {item['line']} 行",
                    "status": "synced",
                }
                for index, item in enumerate(child_links[:12])
            ]
        elif section_links:
            chapter.setdefault("outline_items", [])
    return chapters


def _write_workdraft_markdown(project: dict[str, Any], source_path: Path, markdown: str, stats: dict[str, Any]) -> dict[str, Any]:
    root = _knowledge_root()
    project_dir = root / "06-项目库-Projects" / _safe_filename(project.get("name") or project.get("id") or "文档项目") / "_workdraft"
    project_dir.mkdir(parents=True, exist_ok=True)
    md_path = project_dir / "论文工作稿.md"
    md_path.write_text(markdown, encoding="utf-8")
    return {
        "path": str(md_path),
        "relative_path": _relative_to_vault(md_path),
        "status": "synced",
        "synced_at": stats["generated_at"],
        "generated_from": str(source_path),
        "generated_from_relative": _relative_to_vault(source_path),
        "size_chars": stats["size_chars"],
        "heading_count": stats["heading_count"],
        "paragraph_count": stats["paragraph_count"],
        "table_count": stats["table_count"],
    }


def _read_json_file(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=500, detail=f"Invalid JSON data in {path.name}: {exc}") from exc


def _write_json_file(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _default_agent_templates() -> list[dict[str, Any]]:
    return [
        {
            "id": "leonardo",
            "name": "李奥纳多",
            "role": "架构师 / 开发负责人",
            "team": "ninja_turtles",
            "spawnable": True,
            "capabilities": ["架构设计", "技术方案", "任务拆解", "代码评审"],
            "memory_policy": {"long_term_owner": "leonardo", "project_memory": True, "instance_short_memory": True},
            "default_release_policy": "task_completed",
        },
        {
            "id": "raphael",
            "name": "拉斐尔",
            "role": "后端开发",
            "team": "ninja_turtles",
            "spawnable": True,
            "capabilities": ["后端开发", "API 实现", "数据模型", "接口联调"],
            "memory_policy": {"long_term_owner": "raphael", "project_memory": True, "instance_short_memory": True},
            "default_release_policy": "task_completed",
        },
        {
            "id": "donatello",
            "name": "多纳泰罗",
            "role": "前端开发",
            "team": "ninja_turtles",
            "spawnable": True,
            "capabilities": ["前端开发", "页面实现", "组件集成", "交互联调"],
            "memory_policy": {"long_term_owner": "donatello", "project_memory": True, "instance_short_memory": True},
            "default_release_policy": "task_completed",
        },
        {
            "id": "michelangelo",
            "name": "米开朗基罗",
            "role": "测试工程",
            "team": "ninja_turtles",
            "spawnable": True,
            "capabilities": ["测试工程", "回归验证", "质量检查", "验收清单"],
            "memory_policy": {"long_term_owner": "michelangelo", "project_memory": True, "instance_short_memory": True},
            "default_release_policy": "task_completed",
        },
    ]


def _load_agent_templates() -> list[dict[str, Any]]:
    data = _read_json_file(AGENT_TEMPLATES_FILE, {"templates": _default_agent_templates()})
    templates = data.get("templates") if isinstance(data, dict) else data
    if not isinstance(templates, list) or not templates:
        templates = _default_agent_templates()
    return templates


def _template_by_id() -> dict[str, dict[str, Any]]:
    return {row.get("id"): row for row in _load_agent_templates() if row.get("id")}


def _bootstrap_agent_instances() -> list[dict[str, Any]]:
    return []


def _load_agent_instance_store() -> dict[str, Any]:
    default = {"instances": _bootstrap_agent_instances(), "memory_commits": [], "updated_at": _now_utc()}
    data = _read_json_file(AGENT_INSTANCES_FILE, default)
    if not isinstance(data, dict):
        data = default
    if "instances" not in data or not isinstance(data["instances"], list):
        data["instances"] = _bootstrap_agent_instances()
    if "memory_commits" not in data or not isinstance(data["memory_commits"], list):
        data["memory_commits"] = []
    return data


def _save_agent_instance_store(store: dict[str, Any]) -> dict[str, Any]:
    store["updated_at"] = _now_utc()
    _write_json_file(AGENT_INSTANCES_FILE, store)
    return store


def _enrich_agent_instance(instance: dict[str, Any]) -> dict[str, Any]:
    templates = _template_by_id()
    row = dict(instance)
    template = templates.get(row.get("base_agent_id"), {})
    row.setdefault("base_agent_name", template.get("name", row.get("base_agent_id", "")))
    row.setdefault("role", template.get("role", ""))
    row.setdefault("capabilities", template.get("capabilities", []))
    row.setdefault("type", "elastic")
    row.setdefault("source", "agent-instance-registry")
    return row


def _find_project_task_context(project_id: str, task_id: str, point_id: str = "") -> dict[str, Any]:
    project = project_manager.get_project(project_id) if project_id else None
    task = None
    point = None
    if project:
        for item in project.get("tasks", []):
            if item.get("id") == task_id:
                task = item
                break
        if task:
            for item in task.get("development_points", []):
                if item.get("id") == point_id:
                    point = item
                    break
    return {"project": project, "task": task, "development_point": point}


def _agent_work_type(base_agent_id: str) -> str:
    return {
        "leonardo": "architecture",
        "raphael": "backend",
        "donatello": "frontend",
        "michelangelo": "testing",
    }.get(base_agent_id, "development")


def _instance_dispatch_id(instance: dict[str, Any]) -> str:
    task_id = instance.get("task_id") or "standby"
    return f"elastic-{instance.get('id')}-{task_id}"


def _ensure_instance_workspace(instance: dict[str, Any], template: dict[str, Any], context: dict[str, Any]) -> dict[str, str]:
    instance_id = instance.get("id") or instance.get("instance_id")
    workspace = OPENCLAW_AGENTS_DIR / instance_id / "workspace"
    workspace.mkdir(parents=True, exist_ok=True)
    project = context.get("project") or {}
    task = context.get("task") or {}
    point = context.get("development_point") or {}
    display_name = f"{instance.get('requested_by', '')}-{template.get('name', instance.get('base_agent_id', ''))}"
    identity = (
        f"# {display_name}\n\n"
        f"**Name**: {display_name}\n"
        f"**Role**: {template.get('role', instance.get('role', '弹性执行实例'))}\n"
        f"**Status**: {instance.get('status', 'running')}\n"
        f"**Base Agent**: {instance.get('base_agent_id', '')}\n"
        f"**Requested By**: {instance.get('requested_by', '')}\n"
    )
    agents_md = (
        "# 弹性智能体执行说明\n\n"
        "你是由项目经理按需创建的临时执行分身。\n\n"
        f"- 来源模板：{template.get('name', instance.get('base_agent_id', ''))}\n"
        f"- 职责：{template.get('role', instance.get('role', ''))}\n"
        f"- 能力：{', '.join(template.get('capabilities', []))}\n"
        f"- 请求方：{instance.get('requested_by', '')}\n"
        f"- 项目：{project.get('name', instance.get('project_id', ''))}\n"
        f"- 任务：{task.get('title', instance.get('task_id', ''))}\n"
        f"- 开发要点：{point.get('title', instance.get('development_point_id', '')) or '未指定'}\n\n"
        "执行完成后，需要输出结果摘要、关键决策、阻塞点和下一步建议，并回写项目记忆与基础智能体记忆。\n"
    )
    execution_context = {
        "instance": instance,
        "template": template,
        "project": project,
        "task": task,
        "development_point": point,
        "memory_policy": template.get("memory_policy", {}),
        "release_policy": instance.get("release_policy", "task_completed"),
        "updated_at": _now_utc(),
    }
    current_task = (
        f"# 当前执行任务\n\n"
        f"项目：{project.get('name', instance.get('project_id', ''))}\n\n"
        f"任务：{task.get('title', instance.get('task_id', ''))}\n\n"
        f"开发要点：{point.get('title', instance.get('development_point_id', '')) or '未指定'}\n\n"
        f"调用原因：{instance.get('reason', '')}\n"
    )
    for filename, content in {
        "IDENTITY.md": identity,
        "SOUL.md": identity,
        "AGENTS.md": agents_md,
        "CURRENT_TASK.md": current_task,
    }.items():
        path = workspace / filename
        if filename in {"IDENTITY.md", "SOUL.md", "AGENTS.md"} and path.exists():
            continue
        path.write_text(content, encoding="utf-8")
    _write_json_file(workspace / "EXECUTION_CONTEXT.json", execution_context)
    return {
        "workspace": str(workspace),
        "execution_context": str(workspace / "EXECUTION_CONTEXT.json"),
        "current_task": str(workspace / "CURRENT_TASK.md"),
    }


def _load_dispatch_queue() -> dict[str, Any]:
    default = {"version": "1.0", "sprint": 0, "status": "running", "loop_mode": "auto", "tasks": []}
    data = _read_json_file(NINJA_TURTLES_QUEUE_FILE, default)
    if not isinstance(data, dict):
        data = default
    if "tasks" not in data or not isinstance(data["tasks"], list):
        data["tasks"] = []
    return data


def _save_dispatch_queue(queue: dict[str, Any]) -> dict[str, Any]:
    queue["updated_at"] = _now_utc()
    _write_json_file(NINJA_TURTLES_QUEUE_FILE, queue)
    return queue


def _dispatch_instance_to_openclaw(instance: dict[str, Any], template: dict[str, Any]) -> dict[str, Any]:
    context = _find_project_task_context(instance.get("project_id", ""), instance.get("task_id", ""), instance.get("development_point_id", ""))
    workspace_files = _ensure_instance_workspace(instance, template, context)
    queue = _load_dispatch_queue()
    dispatch_id = _instance_dispatch_id(instance)
    project = context.get("project") or {}
    task = context.get("task") or {}
    point = context.get("development_point") or {}
    dispatch_task = {
        "id": dispatch_id,
        "title": task.get("title") or f"{instance.get('id')} 弹性执行任务",
        "type": _agent_work_type(instance.get("base_agent_id", "")),
        "status": "assigned",
        "priority": task.get("priority", "high"),
        "assignee": instance.get("id"),
        "base_agent_id": instance.get("base_agent_id"),
        "requested_by": instance.get("requested_by"),
        "project_id": instance.get("project_id", ""),
        "project_name": project.get("name", ""),
        "task_id": instance.get("task_id", ""),
        "development_point_id": instance.get("development_point_id", ""),
        "development_point_title": point.get("title", ""),
        "reason": instance.get("reason", ""),
        "release_policy": instance.get("release_policy", "task_completed"),
        "execution_context": workspace_files["execution_context"],
        "created": instance.get("created_at") or _now_utc(),
        "updated_at": _now_utc(),
    }
    existing = next((row for row in queue["tasks"] if row.get("id") == dispatch_id), None)
    if existing:
        existing.update(dispatch_task)
        action = "updated"
    else:
        queue["tasks"].append(dispatch_task)
        action = "queued"
    _save_dispatch_queue(queue)
    return {"action": action, "queue_task": dispatch_task, "workspace": workspace_files}


def _mark_dispatch_instance(instance: dict[str, Any], status: str, result: str = "") -> Optional[dict[str, Any]]:
    dispatch_id = _instance_dispatch_id(instance)
    queue = _load_dispatch_queue()
    row = next((item for item in queue.get("tasks", []) if item.get("id") == dispatch_id), None)
    if not row:
        return None
    row["status"] = status
    row["updated_at"] = _now_utc()
    if result:
        row["result_summary"] = result
    if status in {"done", "completed"}:
        row["completed_at"] = row["updated_at"]
    _save_dispatch_queue(queue)
    return row


def _load_runner_state() -> dict[str, Any]:
    data = _read_json_file(AGENT_RUNNER_FILE, {"runs": [], "updated_at": _now_utc()})
    if not isinstance(data, dict):
        data = {"runs": [], "updated_at": _now_utc()}
    if "runs" not in data or not isinstance(data["runs"], list):
        data["runs"] = []
    return data


def _save_runner_state(state: dict[str, Any]) -> dict[str, Any]:
    state["updated_at"] = _now_utc()
    _write_json_file(AGENT_RUNNER_FILE, state)
    return state


def _runner_executor_status() -> dict[str, Any]:
    configured = os.getenv("ELASTIC_AGENT_EXECUTOR_CMD", "").strip()
    first = shlex.split(configured)[0] if configured else ""
    codex_bin = os.getenv("ELASTIC_AGENT_CODEX_BIN", DEFAULT_CODEX_BIN).strip()
    codex_available = bool(codex_bin and (Path(codex_bin).is_file() or shutil.which(codex_bin)))
    openclaw_bin = os.getenv("ELASTIC_AGENT_OPENCLAW_BIN", DEFAULT_OPENCLAW_BIN).strip()
    openclaw_available = bool(openclaw_bin and (Path(openclaw_bin).is_file() or shutil.which(openclaw_bin)))
    preference = os.getenv("ELASTIC_AGENT_EXECUTOR_PREFERENCE", "openclaw").strip().lower()
    if configured:
        mode = "external-command"
        command = first
        available = bool(first and (Path(first).is_file() or shutil.which(first)))
    elif preference == "openclaw" and openclaw_available:
        mode = "openclaw-agent"
        command = openclaw_bin
        available = True
    elif codex_available:
        mode = "codex-cli"
        command = codex_bin
        available = True
    elif openclaw_available:
        mode = "openclaw-agent"
        command = openclaw_bin
        available = True
    else:
        mode = "claim-only"
        command = codex_bin or openclaw_bin
        available = False
    return {
        "configured": bool(configured or codex_available or openclaw_available),
        "command": command,
        "available": available,
        "mode": mode,
        "preference": preference,
        "codex_bin": codex_bin,
        "codex_available": codex_available,
        "codex_runtime_status": os.getenv("ELASTIC_AGENT_CODEX_STATUS", "unknown"),
        "openclaw_bin": openclaw_bin,
        "openclaw_available": openclaw_available,
    }


def _elastic_runner_status() -> dict[str, Any]:
    data = dict(_ELASTIC_RUNNER_STATE)
    data["task_active"] = bool(_ELASTIC_RUNNER_TASK and not _ELASTIC_RUNNER_TASK.done())
    data["interval_seconds"] = int(os.getenv("ELASTIC_AGENT_RUNNER_INTERVAL", "15"))
    data["execute"] = os.getenv("ELASTIC_AGENT_RUNNER_EXECUTE", "true").lower() != "false"
    return data


async def _elastic_agent_runner_loop(interval_seconds: int, execute: bool, max_seconds: int) -> None:
    assert _ELASTIC_RUNNER_STOP is not None
    _ELASTIC_RUNNER_STATE.update({"enabled": True, "running": True, "last_error": ""})
    while not _ELASTIC_RUNNER_STOP.is_set():
        _ELASTIC_RUNNER_STATE["last_tick_at"] = _now_utc()
        try:
            result = await asyncio.to_thread(
                _runner_tick,
                AgentRunnerTickRequest(execute=execute, max_seconds=max_seconds),
            )
            _ELASTIC_RUNNER_STATE["last_claimed"] = bool(result.get("claimed"))
            _ELASTIC_RUNNER_STATE["last_error"] = ""
            if result.get("claimed"):
                run = result.get("run") or {}
                print(f"[ElasticAgentRunner] claimed {run.get('task_id')} status={run.get('status')}")
        except Exception as exc:
            _ELASTIC_RUNNER_STATE["last_error"] = str(exc)
            print(f"[ElasticAgentRunner] tick failed: {exc}")
        try:
            await asyncio.wait_for(_ELASTIC_RUNNER_STOP.wait(), timeout=interval_seconds)
        except asyncio.TimeoutError:
            pass
    _ELASTIC_RUNNER_STATE["running"] = False


def start_elastic_agent_runner() -> None:
    global _ELASTIC_RUNNER_TASK, _ELASTIC_RUNNER_STOP
    if os.getenv("ELASTIC_AGENT_RUNNER_ENABLED", "true").lower() == "false":
        _ELASTIC_RUNNER_STATE.update({"enabled": False, "running": False, "last_error": "disabled"})
        print("[ElasticAgentRunner] disabled")
        return
    if _ELASTIC_RUNNER_TASK and not _ELASTIC_RUNNER_TASK.done():
        return
    interval_seconds = int(os.getenv("ELASTIC_AGENT_RUNNER_INTERVAL", "15"))
    max_seconds = int(os.getenv("ELASTIC_AGENT_RUNNER_MAX_SECONDS", "1200"))
    execute = os.getenv("ELASTIC_AGENT_RUNNER_EXECUTE", "true").lower() != "false"
    _ELASTIC_RUNNER_STOP = asyncio.Event()
    _ELASTIC_RUNNER_TASK = asyncio.create_task(_elastic_agent_runner_loop(interval_seconds, execute, max_seconds))
    print(f"[ElasticAgentRunner] started interval={interval_seconds}s execute={execute}")


async def stop_elastic_agent_runner() -> None:
    global _ELASTIC_RUNNER_TASK, _ELASTIC_RUNNER_STOP
    if not _ELASTIC_RUNNER_TASK:
        return
    if _ELASTIC_RUNNER_STOP:
        _ELASTIC_RUNNER_STOP.set()
    try:
        await asyncio.wait_for(_ELASTIC_RUNNER_TASK, timeout=5)
    except asyncio.TimeoutError:
        _ELASTIC_RUNNER_TASK.cancel()
    _ELASTIC_RUNNER_TASK = None
    _ELASTIC_RUNNER_STOP = None


def _read_execution_context(queue_task: dict[str, Any]) -> dict[str, Any]:
    context_path = queue_task.get("execution_context")
    if not context_path:
        return {}
    path = Path(context_path)
    if not path.exists():
        return {}
    return _read_json_file(path, {})


def _workspace_for_queue_task(queue_task: dict[str, Any]) -> Path:
    context_path = queue_task.get("execution_context")
    if context_path:
        return Path(context_path).parent
    return OPENCLAW_AGENTS_DIR / queue_task.get("assignee", "unknown") / "workspace"


def _codex_repo_for_queue_task(queue_task: dict[str, Any], context: Optional[dict[str, Any]] = None) -> Path:
    context = context or _read_execution_context(queue_task)
    project = context.get("project") or {}
    task = context.get("task") or {}
    task_context = task.get("context") if isinstance(task.get("context"), dict) else {}
    project_context = project.get("context") if isinstance(project.get("context"), dict) else {}
    candidate_keys = ("repo", "repository", "repo_path", "workspace_path", "code_path")
    for source in (queue_task, task_context, project_context, task, project):
        for key in candidate_keys:
            value = source.get(key) if isinstance(source, dict) else None
            if value:
                path = Path(str(value)).expanduser()
                if path.exists():
                    return path
    return Path(DEFAULT_CODEX_REPO).expanduser()


def _render_runner_prompt(queue_task: dict[str, Any], context: dict[str, Any]) -> str:
    instance = context.get("instance") or {}
    template = context.get("template") or {}
    project = context.get("project") or {}
    task = context.get("task") or {}
    point = context.get("development_point") or {}
    repo = _codex_repo_for_queue_task(queue_task, context)
    capabilities = ", ".join(template.get("capabilities", []) or instance.get("capabilities", []))
    return f"""# OpenClaw 弹性智能体执行任务

你是一个临时弹性执行分身。请按照基础智能体模板执行当前项目任务，并在完成后输出结构化结果摘要。

## 身份
- 实例：{queue_task.get('assignee', instance.get('id', ''))}
- 来源模板：{template.get('name', instance.get('base_agent_id', ''))}
- 职责：{template.get('role', instance.get('role', ''))}
- 能力：{capabilities}
- 请求方：{instance.get('requested_by', queue_task.get('requested_by', ''))}

## 项目任务
- 项目：{project.get('name', queue_task.get('project_name', ''))}
- 项目 ID：{queue_task.get('project_id', '')}
- 代码仓库：{repo}
- 任务：{task.get('title', queue_task.get('title', ''))}
- 任务 ID：{queue_task.get('task_id', '')}
- 开发要点：{point.get('title', queue_task.get('development_point_title', '')) or '未指定'}
- 开发要点 ID：{queue_task.get('development_point_id', '')}
- 调用原因：{queue_task.get('reason', '')}

## 输入上下文
- 执行上下文 JSON：{queue_task.get('execution_context', '')}
- 释放策略：{queue_task.get('release_policy', 'task_completed')}

## 输出要求
最后一条回复必须只包含一个合法 JSON 对象，不要用 Markdown 代码块，也不要以“接下来将做”结束：
{{
  "result_summary": "完成了什么",
  "decisions": ["关键技术决策"],
  "blockers": ["阻塞点"],
  "next_actions": ["下一步建议"],
  "memory_summary": "需要回写到项目和基础智能体长期记忆的摘要",
  "verification": ["实际运行的检查或测试及结果"]
}}

## 执行约束
- 请在上述代码仓库中完成最小必要修改。
- 不要提交 git commit，不要推送。
- 不要回退或覆盖与本任务无关的现有改动。
- 尽量运行能证明修改正确的构建或测试，并在结果中写明。
"""


def _parse_structured_runner_result(value: str) -> Optional[dict[str, Any]]:
    text = str(value or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            return None
        try:
            payload = json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            return None
    if not isinstance(payload, dict):
        return None
    required = {"result_summary", "decisions", "blockers", "next_actions", "memory_summary", "verification"}
    if not required.issubset(payload) or not str(payload.get("result_summary") or "").strip():
        return None
    if not all(isinstance(payload.get(key), list) for key in ("decisions", "blockers", "next_actions", "verification")):
        return None
    return payload


def _update_instance_runtime_status(instance_id: str, status: str, extra: Optional[dict[str, Any]] = None) -> None:
    store = _load_agent_instance_store()
    row = next((item for item in store.get("instances", []) if item.get("id") == instance_id), None)
    if not row:
        return
    row["status"] = status
    row["updated_at"] = _now_utc()
    if extra:
        row.update(extra)
    _save_agent_instance_store(store)


def _claim_next_elastic_task(instance_id: Optional[str] = None) -> Optional[dict[str, Any]]:
    queue = _load_dispatch_queue()
    candidates = []
    for row in queue.get("tasks", []):
        if not (str(row.get("id", "")).startswith("elastic-") or row.get("base_agent_id")):
            continue
        if row.get("status") != "assigned":
            continue
        if instance_id and row.get("assignee") != instance_id:
            continue
        candidates.append(row)
    if not candidates:
        return None
    task = candidates[0]
    executor = _runner_executor_status()
    workspace = _workspace_for_queue_task(task)
    canonical_run = work_run_service.claim(
        dispatch_id=str(task.get("id") or ""),
        project_id=str(task.get("project_id") or ""),
        task_id=str(task.get("task_id") or ""),
        development_point_id=str(task.get("development_point_id") or ""),
        agent_id=str(task.get("assignee") or task.get("base_agent_id") or "elastic-runner"),
        executor=str(executor.get("mode") or "claim-only"),
        workspace=str(workspace),
        input_context={
            "queue_task_id": task.get("id"),
            "project_name": task.get("project_name"),
            "title": task.get("title"),
            "reason": task.get("reason"),
            "release_policy": task.get("release_policy"),
        },
        lease_seconds=int(os.getenv("ELASTIC_AGENT_RUNNER_MAX_SECONDS", "1200")) + 60,
    )
    now = _now_utc()
    task["status"] = "in_progress"
    task["started_at"] = task.get("started_at") or now
    task["updated_at"] = now
    _save_dispatch_queue(queue)
    _update_instance_runtime_status(task.get("assignee", ""), "running", {"runner_status": "claimed"})
    claimed_task = dict(task)
    claimed_task["_canonical_run"] = canonical_run
    return claimed_task


def _run_external_executor(queue_task: dict[str, Any], prompt_path: Path, max_seconds: int) -> dict[str, Any]:
    configured = os.getenv("ELASTIC_AGENT_EXECUTOR_CMD", "").strip()
    executor_status = _runner_executor_status()
    executor_mode = str(executor_status.get("mode") or "claim-only")
    instance_workspace = _workspace_for_queue_task(queue_task)
    context = _read_execution_context(queue_task)
    repo = _codex_repo_for_queue_task(queue_task, context)
    result_path = instance_workspace / "RUNNER_RESULT.md"
    output_path = instance_workspace / "RUNNER_OUTPUT.log"
    env = os.environ.copy()
    env["PATH"] = os.getenv("ELASTIC_AGENT_EXECUTOR_PATH", DEFAULT_CODEX_PATH)

    if not configured and executor_mode == "openclaw-agent":
        openclaw_bin = str(executor_status.get("openclaw_bin") or DEFAULT_OPENCLAW_BIN)
        resolved_openclaw = openclaw_bin if Path(openclaw_bin).is_file() else (shutil.which(openclaw_bin, path=env["PATH"]) or "")
        if not resolved_openclaw:
            return {
                "status": "blocked",
                "failure_code": "openclaw_not_found",
                "reason": "OpenClaw CLI not found; set ELASTIC_AGENT_OPENCLAW_BIN",
                "command": openclaw_bin,
            }
        agent_id = _safe_id(str(queue_task.get("base_agent_id") or queue_task.get("assignee") or "raphael"))
        if agent_id not in _template_by_id():
            return {
                "status": "blocked",
                "failure_code": "invalid_openclaw_agent",
                "reason": f"OpenClaw fallback requires a registered base agent, got {agent_id}",
            }
        session_suffix = _safe_id(str(queue_task.get("id") or "elastic-run"))[-80:]
        prompt = prompt_path.read_text(encoding="utf-8")
        args = [
            resolved_openclaw,
            "agent",
            "--agent",
            agent_id,
            "--session-key",
            f"agent:{agent_id}:elastic-{session_suffix}",
            "--message",
            prompt,
            "--timeout",
            str(max_seconds),
            "--json",
        ]
    elif not configured:
        codex_bin = os.getenv("ELASTIC_AGENT_CODEX_BIN", DEFAULT_CODEX_BIN).strip()
        codex_path = Path(codex_bin)
        resolved_codex = codex_bin if codex_path.is_file() else (shutil.which(codex_bin, path=env["PATH"]) or "")
        if not resolved_codex:
            return {
                "status": "skipped",
                "reason": "Codex CLI not found; set ELASTIC_AGENT_CODEX_BIN or ELASTIC_AGENT_EXECUTOR_CMD",
                "command": codex_bin,
            }
        prompt = prompt_path.read_text(encoding="utf-8")
        args = [
            resolved_codex,
            "--ask-for-approval",
            "never",
            "exec",
            "--cd",
            str(repo),
            "--sandbox",
            "workspace-write",
            "--output-last-message",
            str(result_path),
            prompt,
        ]
    else:
        context_path = queue_task.get("execution_context", "")
        formatted = configured.format(
            prompt=str(prompt_path),
            context=context_path,
            workspace=str(instance_workspace),
            repo=str(repo),
            instance_id=queue_task.get("assignee", ""),
            output=str(result_path),
        )
        args = shlex.split(formatted)

    if not args or not shutil.which(args[0], path=env["PATH"]):
        return {"status": "blocked", "reason": "executor command not found", "command": args[0] if args else ""}
    try:
        result = subprocess.run(args, cwd=str(repo), text=True, capture_output=True, timeout=max_seconds, env=env)
    except subprocess.TimeoutExpired as exc:
        timeout_stdout = exc.stdout.decode() if isinstance(exc.stdout, bytes) else (exc.stdout or "")
        timeout_stderr = exc.stderr.decode() if isinstance(exc.stderr, bytes) else (exc.stderr or "")
        output_path.write_text(
            f"$ {' '.join(args)}\n\n[timeout]\n{max_seconds}s\n\n[stdout]\n{timeout_stdout}\n\n[stderr]\n{timeout_stderr}\n",
            encoding="utf-8",
        )
        return {
            "status": "failed",
            "failure_code": "executor_timeout",
            "reason": f"executor exceeded {max_seconds}s timeout",
            "output": str(output_path),
            "repo": str(repo),
            "workspace": str(instance_workspace),
        }
    output_path.write_text(
        f"$ {' '.join(args)}\n\n[stdout]\n{result.stdout}\n\n[stderr]\n{result.stderr}\n",
        encoding="utf-8",
    )
    final_message = result_path.read_text(encoding="utf-8") if result_path.exists() else ""
    openclaw_status = ""
    structured_result: Optional[dict[str, Any]] = None
    if executor_mode == "openclaw-agent":
        try:
            payload = json.loads(result.stdout)
            openclaw_status = str(payload.get("status") or "")
            result_payload = payload.get("result") if isinstance(payload.get("result"), dict) else {}
            payloads = result_payload.get("payloads") if isinstance(result_payload.get("payloads"), list) else []
            messages = [str(item.get("text") or "") for item in payloads if isinstance(item, dict) and item.get("text")]
            if not messages:
                meta = result_payload.get("meta") if isinstance(result_payload.get("meta"), dict) else {}
                visible = meta.get("finalAssistantVisibleText") or meta.get("finalAssistantRawText")
                if visible:
                    messages = [str(visible)]
            final_message = "\n\n".join(messages).strip()
        except (TypeError, json.JSONDecodeError):
            final_message = result.stdout.strip()[-4000:]
        structured_result = _parse_structured_runner_result(final_message)
        if structured_result:
            final_message = json.dumps(structured_result, ensure_ascii=False, indent=2)
        if final_message:
            result_path.write_text(final_message, encoding="utf-8")
    if not final_message.strip() and result.returncode != 0:
        stderr_tail = result.stderr.strip()[-2000:]
        stdout_tail = result.stdout.strip()[-1000:]
        final_message = "\n".join(
            part for part in [
                f"Codex CLI 执行失败，退出码：{result.returncode}",
                f"stderr 摘要：\n{stderr_tail}" if stderr_tail else "",
                f"stdout 摘要：\n{stdout_tail}" if stdout_tail else "",
            ]
            if part
        )
    completed = result.returncode == 0 and (
        executor_mode != "openclaw-agent"
        or (openclaw_status in {"", "ok"} and structured_result is not None)
    )
    failure_code = ""
    failure_reason = ""
    if result.returncode != 0:
        failure_code = "executor_nonzero_exit"
        failure_reason = f"executor exited with code {result.returncode}"
    elif executor_mode == "openclaw-agent" and openclaw_status not in {"", "ok"}:
        failure_code = "openclaw_run_failed"
        failure_reason = f"OpenClaw returned status {openclaw_status}"
    elif executor_mode == "openclaw-agent" and structured_result is None:
        failure_code = "incomplete_executor_result"
        failure_reason = "executor final response did not satisfy the structured completion contract"
    return {
        "status": "completed" if completed else "failed",
        "executor": executor_mode,
        "failure_code": failure_code or None,
        "reason": failure_reason,
        "returncode": result.returncode,
        "output": str(output_path),
        "result": str(result_path) if result_path.exists() else "",
        "summary": final_message[:2000],
        "structured_result": structured_result or {},
        "repo": str(repo),
        "workspace": str(instance_workspace),
    }


def _runner_tick(req: AgentRunnerTickRequest) -> dict[str, Any]:
    try:
        task = _claim_next_elastic_task(req.instance_id)
    except WorkRunLeaseConflict as exc:
        return {
            "claimed": False,
            "reason": "lease_conflict",
            "detail": str(exc),
            "executor": _runner_executor_status(),
        }
    if not task:
        return {"claimed": False, "reason": "no assigned elastic task", "executor": _runner_executor_status()}
    canonical_run = task.pop("_canonical_run")
    context = _read_execution_context(task)
    workspace = _workspace_for_queue_task(task)
    workspace.mkdir(parents=True, exist_ok=True)
    prompt_path = workspace / "RUNNER_PROMPT.md"
    prompt_path.write_text(_render_runner_prompt(task, context), encoding="utf-8")
    canonical_run = work_run_service.transition(
        canonical_run["id"],
        "claimed",
        actor=str(task.get("assignee") or "elastic-runner"),
        detail="Runner prompt and workspace prepared",
        prompt_path=str(prompt_path),
        workspace=str(workspace),
    )
    run = {
        **canonical_run,
        "task_id": canonical_run.get("task_id") or task.get("task_id"),
        "dispatch_id": task.get("id"),
        "instance_id": task.get("assignee"),
        "execute_requested": req.execute,
        "prompt": str(prompt_path),
        "execution_context": task.get("execution_context", ""),
    }
    if req.execute:
        canonical_run = work_run_service.transition(
            canonical_run["id"],
            "running",
            actor=str(task.get("assignee") or "elastic-runner"),
            detail="External executor started",
            lease_seconds=req.max_seconds + 60,
        )
        execution = _run_external_executor(task, prompt_path, req.max_seconds)
        run["execution"] = execution
        run["status"] = execution.get("status", "unknown")
        if execution.get("status") == "completed":
            canonical_run = work_run_service.transition(
                canonical_run["id"],
                "review",
                actor=str(task.get("assignee") or "elastic-runner"),
                detail="Executor completed; awaiting independent verification",
                result_summary=execution.get("summary", ""),
                execution_result=execution,
            )
            for artifact_type, title, uri in (
                ("runner-output", "Executor stdout and stderr", execution.get("output", "")),
                ("runner-result", "Executor result summary", execution.get("result", "")),
            ):
                if uri:
                    work_run_service.add_artifact(
                        canonical_run["id"],
                        artifact_type=artifact_type,
                        title=title,
                        uri=str(uri),
                        metadata={"executor": _runner_executor_status().get("mode")},
                    )
            task["status"] = "review"
            task["updated_at"] = _now_utc()
            task["result_summary"] = execution.get("summary", "")
            _update_instance_runtime_status(
                task.get("assignee", ""),
                "review",
                {"runner_status": "completed", "last_result_summary": execution.get("summary", "")},
            )
            queue = _load_dispatch_queue()
            existing = next((row for row in queue.get("tasks", []) if row.get("id") == task.get("id")), None)
            if existing:
                existing.update(task)
                _save_dispatch_queue(queue)
        elif execution.get("status") in {"failed", "blocked", "skipped"}:
            canonical_status = "failed" if execution.get("status") == "failed" else "blocked"
            failure_detail = execution.get("summary") or execution.get("reason", "")
            canonical_run = work_run_service.transition(
                canonical_run["id"],
                canonical_status,
                actor=str(task.get("assignee") or "elastic-runner"),
                detail="External executor did not produce a reviewable result",
                result_summary=failure_detail,
                failure_code=execution.get("failure_code") or f"executor_{execution.get('status')}",
                failure_detail=failure_detail,
                execution_result=execution,
            )
            if execution.get("output"):
                work_run_service.add_artifact(
                    canonical_run["id"],
                    artifact_type="runner-output",
                    title="Failed executor output",
                    uri=str(execution["output"]),
                    metadata={"status": execution.get("status")},
                )
            task["status"] = "blocked"
            task["updated_at"] = _now_utc()
            task["result_summary"] = failure_detail
            _update_instance_runtime_status(
                task.get("assignee", ""),
                "blocked",
                {"runner_status": execution.get("status"), "last_result_summary": task["result_summary"]},
            )
            queue = _load_dispatch_queue()
            existing = next((row for row in queue.get("tasks", []) if row.get("id") == task.get("id")), None)
            if existing:
                existing.update(task)
                _save_dispatch_queue(queue)
        run = {
            **work_run_service.get(canonical_run["id"]),
            "dispatch_id": task.get("id"),
            "instance_id": task.get("assignee"),
            "execute_requested": req.execute,
            "prompt": str(prompt_path),
            "execution_context": task.get("execution_context", ""),
            "execution": execution,
        }
    runner_status_path = workspace / "RUNNER_STATUS.json"
    runner_status_path.write_text(json.dumps(run, ensure_ascii=False, indent=2), encoding="utf-8")
    state = _load_runner_state()
    state.setdefault("runs", []).append(run)
    state["runs"] = state["runs"][-200:]
    _save_runner_state(state)
    return {"claimed": True, "queue_task": task, "run": run, "executor": _runner_executor_status()}


@router.get("/projects")
def list_projects(
    project_type: Optional[str] = Query(None, description="Legacy project template filter: software or document"),
    enabled_module: Optional[str] = Query(None, description="Filter by composable module key"),
):
    projects = project_manager.list_projects()
    if project_type:
        normalized_type = project_type.strip().lower()
        if normalized_type not in {"software", "document"}:
            raise HTTPException(status_code=400, detail="project_type must be software or document")
        projects = [
            project for project in projects
            if str(project.get("project_type") or project.get("type") or "software").lower() == normalized_type
        ]
    if enabled_module:
        from services.project_composition import PROJECT_MODULES
        if enabled_module not in PROJECT_MODULES:
            raise HTTPException(status_code=400, detail="enabled_module is invalid")
        projects = [
            project for project in projects
            if enabled_module in (project.get("enabled_modules") or [])
        ]
    return {"projects": projects, "total": len(projects)}


@router.post("/projects", status_code=201)
def create_project(req: ProjectCreate):
    return project_manager.create_project(_model_dict(req))


@router.get("/projects/{project_id}")
def get_project(project_id: str):
    project = project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@router.put("/projects/{project_id}")
def update_project(project_id: str, req: ProjectUpdate):
    project = project_manager.update_project(project_id, _model_dict(req, exclude_unset=True))
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@router.get("/projects/{project_id}/document-workdraft")
def get_project_document_workdraft(project_id: str):
    project = project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    spec = project.get("document_spec") if isinstance(project.get("document_spec"), dict) else {}
    return {
        "project_id": project_id,
        "source_word": spec.get("source_word", {}),
        "working_markdown": spec.get("working_markdown", {}),
        "section_links": spec.get("section_links", []),
        "sync_status": spec.get("sync_status", {}),
    }


@router.post("/projects/{project_id}/document-workdraft/sync")
def sync_project_document_workdraft(project_id: str, req: DocumentWorkdraftSyncRequest):
    project = project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if str(project.get("project_type") or project.get("type") or "").lower() != "document":
        raise HTTPException(status_code=400, detail="Only document projects can create a workdraft")

    source_path = _find_word_source(project, req.source_word_path)
    markdown, stats = _docx_to_markdown(source_path, project)
    working_markdown = _write_workdraft_markdown(project, source_path, markdown, stats)

    spec = dict(project.get("document_spec") or {})
    chapters = [dict(item) for item in spec.get("chapters", []) if isinstance(item, dict)]
    section_links = _extract_markdown_headings(markdown, chapters)
    spec.update({
        "source_word": {
            "path": str(source_path),
            "relative_path": _relative_to_vault(source_path),
            "title": source_path.stem,
            "file_name": source_path.name,
            "size_bytes": source_path.stat().st_size,
            "mtime": datetime.fromtimestamp(source_path.stat().st_mtime, timezone.utc).isoformat(),
            "knowledge_node_id": _relative_to_vault(source_path),
        },
        "working_markdown": working_markdown,
        "section_links": section_links,
        "sync_status": {
            "status": "synced",
            "message": "Word 原文已转换为 Markdown 工作稿，并生成章节绑定",
            "agent_id": req.agent_id,
            "synced_at": stats["generated_at"],
            "heading_count": stats["heading_count"],
            "section_link_count": len(section_links),
        },
    })
    spec["chapters"] = _chapters_with_workdraft_outline(project, section_links)

    updated = project_manager.update_project(project_id, {"document_spec": spec})
    if not updated:
        raise HTTPException(status_code=404, detail="Project not found")
    return {
        "project": updated,
        "source_word": spec["source_word"],
        "working_markdown": working_markdown,
        "section_links": section_links,
        "markdown_preview": markdown[:2000],
    }


@router.delete("/projects/{project_id}")
def delete_project(project_id: str):
    deleted = project_manager.delete_project(project_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Project not found")
    return {"project_id": project_id, "deleted": True}


@router.get("/projects/{project_id}/design-doc")
def get_project_design_doc(project_id: str):
    design_doc = project_manager.get_design_doc(project_id)
    if not design_doc:
        raise HTTPException(status_code=404, detail="Project not found")
    return design_doc


@router.put("/projects/{project_id}/design-doc")
def update_project_design_doc(project_id: str, req: DesignDocumentUpdate):
    payload = _model_dict(req, exclude_unset=True)
    agent_id = payload.get("author_agent") or "project-manager"
    design_doc = project_manager.update_design_doc(project_id, payload, agent_id)
    if not design_doc:
        raise HTTPException(status_code=404, detail="Project not found")
    return design_doc


@router.post("/projects/{project_id}/design-doc/revise", status_code=201)
def revise_project_design_doc(project_id: str, req: DesignDocumentRevise):
    result = project_manager.revise_design_doc(project_id, _model_dict(req))
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result


@router.post("/projects/{project_id}/design-doc/approve")
def approve_project_design_doc(project_id: str, req: ApproveDesignDocumentRequest):
    result = project_manager.approve_design_doc(project_id, req.agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result


@router.get("/projects/{project_id}/agent-context")
def get_project_agent_context(project_id: str, agent_id: Optional[str] = Query(None)):
    context = project_manager.get_agent_context(project_id, data_manager.get_agents(), agent_id)
    if not context:
        raise HTTPException(status_code=404, detail="Project not found")
    return context


@router.get("/projects/{project_id}/tasks")
def list_project_tasks(project_id: str):
    tasks = project_manager.list_tasks(project_id)
    if tasks is None:
        raise HTTPException(status_code=404, detail="Project not found")
    return {"project_id": project_id, "tasks": tasks, "total": len(tasks)}


@router.post("/projects/{project_id}/tasks", status_code=201)
def create_project_task(project_id: str, req: TaskCreate):
    task = project_manager.add_task(project_id, _model_dict(req))
    if not task:
        raise HTTPException(status_code=404, detail="Project not found")
    # 同步到 V2 任务列表
    project = project_manager.get_project(project_id)
    if project:
        sync_v3_task_to_v2(
            task,
            project_id=project_id,
            project_name=project.get("name", ""),
            project_type=project.get("project_type", "software"),
        )
    return task


@router.delete("/projects/{project_id}/tasks/{task_id}")
def delete_project_task(project_id: str, task_id: str):
    deleted = project_manager.delete_task(project_id, task_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Task not found")

    try:
        from database import get_session
        from services.task_service import TaskService

        db = get_session()
        try:
            TaskService(db).delete_task_by_task_id(f"v3-{project_id}-{task_id}")
        finally:
            db.close()
    except Exception as exc:
        import logging
        logging.getLogger(__name__).warning("V3 project task delete sync failed: %s", exc)

    try:
        queue = _load_dispatch_queue()
        before = len(queue.get("tasks", []))
        queue["tasks"] = [
            row for row in queue.get("tasks", [])
            if not (row.get("project_id") == project_id and row.get("task_id") == task_id)
        ]
        if len(queue["tasks"]) != before:
            _save_dispatch_queue(queue)
    except Exception as exc:
        import logging
        logging.getLogger(__name__).warning("V3 project task queue cleanup failed: %s", exc)

    return {"project_id": project_id, "task_id": task_id, "deleted": True}


@router.put("/tasks/{task_id}")
def update_task(task_id: str, req: TaskUpdate):
    task = project_manager.update_task(task_id, _model_dict(req, exclude_unset=True))
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    # 同步状态/进度变更到 V2 任务列表
    try:
        from project_manager import _as_list
        from services.project_task_sync import map_v3_status, map_v3_priority
        from database import get_session
        from services.task_service import TaskService

        # 查找所属项目
        all_projects = project_manager.list_projects()
        for p in all_projects:
            for t in p.get("tasks", []):
                if t.get("id") == task_id:
                    v2_task_id = f"v3-{p['id']}-{task_id}"
                    db = get_session()
                    try:
                        svc = TaskService(db)
                        updates = {}
                        if req.status is not None:
                            updates["status"] = map_v3_status(req.status)
                        if req.progress is not None:
                            updates["progress"] = int(req.progress)
                            if req.status in ("done", "completed"):
                                updates["progress"] = 100
                        if req.priority is not None:
                            updates["priority"] = map_v3_priority(req.priority)
                        if req.title is not None:
                            updates["title"] = req.title
                        if req.description is not None:
                            updates["description"] = req.description
                        if updates:
                            svc.update_task_by_task_id(v2_task_id, updates, changed_by="system-sync")
                    finally:
                        db.close()
                    break
    except Exception as exc:
        import logging
        logging.getLogger(__name__).error(f"V3→V2 状态同步失败: {exc}")
    return task


@router.post("/tasks/{task_id}/assign")
def assign_task(task_id: str, req: AssignRequest):
    task = project_manager.assign_task(task_id, req.assignee_agent)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return task


@router.post("/projects/sync-all-tasks")
def sync_all_project_tasks_to_v2():
    """一键批量同步：将所有 V3 项目任务同步到 V2 任务列表"""
    results = []
    projects = project_manager.list_projects()
    for project in projects:
        pid = project.get("id", "")
        pname = project.get("name", "")
        ptype = project.get("project_type", "software")
        tasks = project.get("tasks", [])
        synced = 0
        failed = 0
        for task in tasks:
            result = sync_v3_task_to_v2(task, pid, pname, ptype)
            if result:
                synced += 1
            else:
                failed += 1
        results.append({
            "project_id": pid,
            "project_name": pname,
            "synced": synced,
            "failed": failed,
            "total": len(tasks),
        })
    return {"results": results}


@router.get("/tasks/{task_id}/points")
def list_task_points(task_id: str):
    points = project_manager.list_points(task_id)
    if points is None:
        raise HTTPException(status_code=404, detail="Task not found")
    return {"task_id": task_id, "points": points, "total": len(points)}


@router.post("/tasks/{task_id}/points", status_code=201)
def create_task_point(task_id: str, req: DevelopmentPointCreate):
    point = project_manager.add_point(task_id, _model_dict(req))
    if not point:
        raise HTTPException(status_code=404, detail="Task not found")
    return point


@router.put("/points/{point_id}")
def update_development_point(point_id: str, req: DevelopmentPointUpdate):
    point = project_manager.update_point(point_id, _model_dict(req, exclude_unset=True))
    if not point:
        raise HTTPException(status_code=404, detail="Development point not found")
    return point


@router.post("/points/{point_id}/complete")
def complete_development_point(point_id: str, req: CompletePointRequest):
    result = project_manager.complete_point(point_id, req.agent_id, req.completion_evidence, req.result_summary)
    if not result:
        raise HTTPException(status_code=404, detail="Development point not found")
    return result




def _transition_point(point_id: str, action: str, req: PointTransitionRequest):
    result = project_manager.transition_point(
        point_id,
        action,
        req.agent_id,
        req.reason,
        req.completion_evidence,
        req.result_summary,
    )
    if not result:
        raise HTTPException(status_code=404, detail="Development point not found")
    return result


@router.post("/points/{point_id}/claim")
def claim_development_point(point_id: str, req: PointTransitionRequest):
    return _transition_point(point_id, "claim", req)


@router.post("/points/{point_id}/release")
def release_development_point(point_id: str, req: PointTransitionRequest):
    return _transition_point(point_id, "release", req)


@router.post("/points/{point_id}/block")
def block_development_point(point_id: str, req: PointTransitionRequest):
    return _transition_point(point_id, "block", req)


@router.post("/points/{point_id}/submit-review")
def submit_development_point_review(point_id: str, req: PointTransitionRequest):
    return _transition_point(point_id, "submit_review", req)

@router.post("/tasks/{task_id}/logs", status_code=201)
def create_task_log(task_id: str, req: LogCreate):
    project_id = None
    for project in project_manager.list_projects():
        if any(task.get("id") == task_id for task in project.get("tasks", [])):
            project_id = project["id"]
            break
    if not project_id:
        raise HTTPException(status_code=404, detail="Task not found")
    log = project_manager.add_log(project_id, task_id, req.agent_id, req.action, req.content)
    return log


@router.get("/projects/{project_id}/logs")
def list_project_logs(project_id: str, limit: int = Query(50, ge=1, le=200)):
    if not project_manager.get_project(project_id):
        raise HTTPException(status_code=404, detail="Project not found")
    return {"project_id": project_id, "logs": project_manager.list_logs(project_id, limit)}


@router.get("/projects/{project_id}/conversation")
def list_project_conversation(project_id: str, limit: int = Query(80, ge=1, le=200)):
    messages = project_manager.list_conversation(project_id, limit)
    if messages is None:
        raise HTTPException(status_code=404, detail="Project not found")
    return {"project_id": project_id, "messages": messages, "total": len(messages)}


@router.get("/projects/{project_id}/chat-context")
def get_project_chat_context(project_id: str):
    context = project_manager.build_project_chat_context(project_id)
    if not context:
        raise HTTPException(status_code=404, detail="Project not found")
    return context


@router.post("/projects/{project_id}/chat", status_code=201)
def create_project_chat_message(project_id: str, req: ProjectChatRequest):
    payload = _model_dict(req)
    result = project_manager.add_conversation_message(project_id, payload)
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    context = project_manager.build_project_chat_context(project_id)
    project_type = (context or {}).get("project", {}).get("project_type", "software")
    if project_type == "document":
        next_actions = [
            "根据目录结构拆分章节写作任务",
            "补充章节主要内容和图片/图表计划",
            "把引用资料绑定到章节或写作任务",
        ]
    else:
        next_actions = [
            "根据设计文档拆分开发任务",
            "指派后端、前端、测试智能体并行推进",
            "把接口、数据结构和验收标准写入开发要点",
        ]
    return {
        **result,
        "context": context,
        "suggested_next_actions": next_actions,
        "dispatch_status": "recorded",
    }


@router.post("/projects/{project_id}/agent-actions", status_code=201)
def create_project_agent_action(project_id: str, req: ProjectAgentActionRequest):
    project = project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    project_type = project.get("project_type") or project.get("type") or project.get("context", {}).get("project_type") or "software"
    instruction = req.instruction or req.payload.get("instruction") or req.task_title or req.action_type
    chat = project_manager.add_conversation_message(project_id, {
        "agent_id": req.agent_id,
        "role": "system",
        "intent": "agent_action",
        "message": instruction,
        "task_id": req.task_id,
        "chapter_id": req.chapter_id,
    })
    created_task = None
    should_create_task = req.action_type in {"create_task", "assign_task", "writing_task", "development_task"} or bool(req.task_title)
    if should_create_task:
        task_type = req.task_type or req.payload.get("task_type") or ("writing" if project_type == "document" else "development")
        title = req.task_title or req.payload.get("title") or instruction
        created_task = project_manager.add_task(project_id, {
            "title": title,
            "description": instruction,
            "type": task_type,
            "assignee_agent": req.assignee_agent or req.agent_id,
            "assignee_agent_id": req.assignee_agent or req.agent_id,
            "status": "todo",
            "priority": req.payload.get("priority", "medium"),
            "context": {
                "created_from_agent_action": req.action_type,
                "chapter_id": req.chapter_id,
                "project_type": project_type,
            },
            "development_points": [{
                "title": req.payload.get("point_title") or (title + ("写作要点" if task_type == "writing" else "开发要点")),
                "status": "todo",
                "assigned_agent": req.assignee_agent or req.agent_id,
            }],
        })
    log = project_manager.add_log(project_id, req.task_id or (created_task or {}).get("id"), req.agent_id, "project_agent_action", instruction)
    return {
        "project_id": project_id,
        "action_type": req.action_type,
        "agent_id": req.agent_id,
        "chat": chat,
        "created_task": created_task,
        "log": log,
        "dispatch_status": "recorded",
    }


@router.put("/projects/{project_id}/software-spec")
def update_project_software_spec(project_id: str, req: SoftwareSpecUpdate):
    payload = _model_dict(req, exclude_unset=True)
    agent_id = payload.pop("agent_id", "project-manager")
    result = project_manager.update_software_spec(project_id, payload, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result


@router.post("/projects/{project_id}/document-sections", status_code=201)
def create_document_section(project_id: str, req: DocumentSectionCreate):
    payload = _model_dict(req, exclude_unset=True)
    agent_id = payload.pop("agent_id", "project-manager")
    result = project_manager.add_document_section(project_id, payload, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result


@router.put("/projects/{project_id}/document-sections/{section_id}")
def update_document_section(project_id: str, section_id: str, req: DocumentSectionUpdate):
    payload = _model_dict(req, exclude_unset=True)
    agent_id = payload.pop("agent_id", "project-manager")
    result = project_manager.update_document_section(project_id, section_id, payload, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project or section not found")
    return result


@router.delete("/projects/{project_id}/document-sections/{section_id}")
def delete_document_section(project_id: str, section_id: str, agent_id: str = Query("project-manager")):
    result = project_manager.delete_document_section(project_id, section_id, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project or section not found")
    return result


@router.post("/projects/{project_id}/document-assets", status_code=201)
def create_document_asset(project_id: str, req: DocumentAssetCreate):
    payload = _model_dict(req, exclude_unset=True)
    agent_id = payload.pop("agent_id", "project-manager")
    result = project_manager.add_document_asset(project_id, payload, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result


@router.put("/projects/{project_id}/document-assets/{asset_id}")
def update_document_asset(project_id: str, asset_id: str, req: DocumentAssetUpdate):
    payload = _model_dict(req, exclude_unset=True)
    agent_id = payload.pop("agent_id", "project-manager")
    result = project_manager.update_document_asset(project_id, asset_id, payload, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project or asset not found")
    return result


@router.delete("/projects/{project_id}/document-assets/{asset_id}")
def delete_document_asset(project_id: str, asset_id: str, agent_id: str = Query("project-manager")):
    result = project_manager.delete_document_asset(project_id, asset_id, agent_id)
    if not result:
        raise HTTPException(status_code=404, detail="Project or asset not found")
    return result


@router.get("/projects/{project_id}/iteration-context")
def get_iteration_context(project_id: str):
    context = project_manager.get_iteration_context(project_id, data_manager.get_agents())
    if not context:
        raise HTTPException(status_code=404, detail="Project not found")
    context["knowledge_context"] = knowledge_manager.project_context(context["project"])
    return context


@router.get("/projects/{project_id}/knowledge-context")
def get_project_knowledge_context(project_id: str):
    project = project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return knowledge_manager.project_context(project)


@router.post("/projects/{project_id}/knowledge-links", status_code=201)
def add_project_knowledge_link(project_id: str, req: KnowledgeLinkRequest):
    result = project_manager.add_knowledge_link("project", project_id, _model_dict(req))
    if not result:
        raise HTTPException(status_code=404, detail="Project not found or invalid knowledge link")
    project = project_manager.get_project(project_id)
    result["knowledge_context"] = knowledge_manager.project_context(project) if project else None
    return result


@router.post("/tasks/{task_id}/knowledge-links", status_code=201)
def add_task_knowledge_link(task_id: str, req: KnowledgeLinkRequest):
    result = project_manager.add_knowledge_link("task", task_id, _model_dict(req))
    if not result:
        raise HTTPException(status_code=404, detail="Task not found or invalid knowledge link")
    return result


@router.post("/points/{point_id}/knowledge-links", status_code=201)
def add_point_knowledge_link(point_id: str, req: KnowledgeLinkRequest):
    result = project_manager.add_knowledge_link("point", point_id, _model_dict(req))
    if not result:
        raise HTTPException(status_code=404, detail="Development point not found or invalid knowledge link")
    return result


@router.delete("/knowledge-links")
def delete_knowledge_link(
    target_type: str = Query(..., pattern="^(project|task|point)$"),
    target_id: str = Query(..., min_length=1),
    node_id: str = Query(..., min_length=1),
    removed_by: str = Query("project-manager"),
):
    result = project_manager.remove_knowledge_link(target_type, target_id, node_id, removed_by)
    if not result:
        raise HTTPException(status_code=404, detail="Knowledge link not found")
    if target_type == "project":
        project = project_manager.get_project(target_id)
        result["knowledge_context"] = knowledge_manager.project_context(project) if project else None
    return result


@router.post("/projects/{project_id}/decompose", status_code=201)
def decompose_project(project_id: str, req: DecomposeRequest):
    result = project_manager.decompose_project(project_id, _model_dict(req))
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result

# Agent status projection backed by V3 project tasks.
def _now_utc():
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _merge_agent_sources(openclaw_agents, local_agents):
    merged = {}
    for agent in local_agents:
        aid = agent.get("id")
        if not aid:
            continue
        row = dict(agent)
        row["_dashboard_source"] = "local-data-manager"
        merged[aid] = row

    # Fields that OpenClaw runtime can provide (status, last_seen)
    # All other fields (name, role, team, memory, config, etc.) stay from local agents.json
    runtime_fields = {"status", "current_task", "last_seen", "last_heartbeat",
                      "cpu_usage", "memory_usage", "task_queue_len", "metadata",
                      "updated_at", "agent_name"}

    for agent in openclaw_agents:
        aid = agent.get("id")
        if not aid:
            continue
        if aid in merged:
            # Local data exists → only overlay runtime fields, keep local profile data
            for key in runtime_fields:
                if key in agent and agent[key] is not None:
                    merged[aid][key] = agent[key]
            merged[aid]["_dashboard_source"] = "openclaw-workspace+local-data-manager"
        else:
            # No local data → use OpenClaw data as-is (new/unknown agents)
            row = dict(agent)
            row["_dashboard_source"] = "openclaw-workspace"
            merged[aid] = row
    return list(merged.values())


def _load_agent_sources():
    observed_at = _now_utc()
    local_agents = [
        agent for agent in data_manager.get_agents()
        if agent.get("enabled", True)
    ]
    try:
        openclaw_agents = openclaw_integration.sync_agents()
    except Exception:
        openclaw_agents = []

    if openclaw_agents:
        return _merge_agent_sources(openclaw_agents, local_agents), "mixed", observed_at

    agents = []
    for agent in local_agents:
        row = dict(agent)
        row["_dashboard_source"] = "local-data-manager"
        agents.append(row)
    return agents, "local-data-manager", observed_at


def _agent_rows(agents: Optional[list[dict[str, Any]]] = None):
    projects = project_manager.list_projects()
    agents = agents if agents is not None else data_manager.get_agents()
    rows = []
    done = {"done", "completed"}
    for agent in agents:
        aid = agent.get("id", "")
        row = {
            "agent_id": aid,
            "agent_name": agent.get("name", aid),
            "role": agent.get("role", ""),
            "team": agent.get("team", ""),
            "status": agent.get("status", "idle"),
            "legacy_current_task": agent.get("current_task", ""),
            "current_project_id": None,
            "current_project_name": None,
            "current_task_id": None,
            "current_task_title": None,
            "current_development_point_id": None,
            "current_development_point_title": None,
            "task_progress": 0,
            "project_progress": 0,
            "source": "projects-v3",
        }
        for project in projects:
            for task in project.get("tasks", []):
                task_agent = task.get("assignee_agent")
                if task_agent == aid and task.get("status") not in done:
                    row.update({"current_project_id": project.get("id"), "current_project_name": project.get("name"), "current_task_id": task.get("id"), "current_task_title": task.get("title"), "task_progress": task.get("progress", 0), "project_progress": project.get("progress", 0), "status": "busy"})
                for point in task.get("development_points", []):
                    point_agent = point.get("assigned_agent") or task_agent
                    if point_agent == aid and point.get("status") not in done:
                        row.update({"current_project_id": project.get("id"), "current_project_name": project.get("name"), "current_task_id": task.get("id"), "current_task_title": task.get("title"), "current_development_point_id": point.get("id"), "current_development_point_title": point.get("title"), "task_progress": task.get("progress", 0), "project_progress": project.get("progress", 0), "status": "busy"})
                        break
                if row["current_development_point_id"]:
                    break
            if row["current_development_point_id"]:
                break
        rows.append(row)
    return rows



@router.get("/skills")
def list_skills(query: str = "", category: str = "", source: str = "", status: str = ""):
    rows = skill_manager.list_skills(query=query, category=category, source=source, status=status)
    return {
        "skills": rows,
        "total": len(rows),
        "categories": sorted({row.get("category", "") for row in rows if row.get("category")}),
        "sources": sorted({row.get("source", "") for row in rows if row.get("source")}),
        "source": "skill-registry",
        "updated_at": _now_utc(),
    }


@router.get("/skills/{skill_id}")
def get_skill(skill_id: str):
    skill = skill_manager.get_skill(skill_id)
    if not skill:
        raise HTTPException(status_code=404, detail="Skill not found")
    return skill


@router.get("/agents/{agent_id}/skills")
def get_agent_skills(agent_id: str):
    rows = skill_manager.get_agent_skills(agent_id)
    return {"agent_id": agent_id, "skills": rows, "total": len(rows)}


@router.put("/agents/{agent_id}/skills")
def update_agent_skills(agent_id: str, req: AgentSkillsUpdate):
    rows = skill_manager.set_agent_skills(agent_id, req.skill_ids)
    return {"agent_id": agent_id, "skills": rows, "total": len(rows)}


@router.get("/agents/templates")
def list_agent_templates():
    templates = _load_agent_templates()
    return {
        "templates": templates,
        "total": len(templates),
        "source": "agent-template-registry",
        "updated_at": _now_utc(),
    }


@router.get("/agents/instances")
def list_agent_instances(status: str = ""):
    store = _load_agent_instance_store()
    rows = [_enrich_agent_instance(row) for row in store.get("instances", [])]
    if status:
        wanted = {item.strip() for item in status.split(",") if item.strip()}
        rows = [row for row in rows if row.get("status") in wanted]
    return {
        "instances": rows,
        "total": len(rows),
        "memory_commits": store.get("memory_commits", []),
        "source": "agent-instance-registry",
        "updated_at": store.get("updated_at") or _now_utc(),
    }


@router.get("/agents/dispatch-queue")
def get_agent_dispatch_queue():
    queue = _load_dispatch_queue()
    elastic_tasks = [
        row for row in queue.get("tasks", [])
        if str(row.get("id", "")).startswith("elastic-") or row.get("base_agent_id")
    ]
    return {
        "queue": queue,
        "elastic_tasks": elastic_tasks,
        "total": len(elastic_tasks),
        "source": "openclaw-dev-loop-queue",
        "updated_at": queue.get("updated_at") or _now_utc(),
    }


@router.get("/agents/runner/status")
def get_agent_runner_status():
    state = _load_runner_state()
    canonical = work_run_service.summary(limit=200)
    queue = _load_dispatch_queue()
    assigned = [
        row for row in queue.get("tasks", [])
        if (str(row.get("id", "")).startswith("elastic-") or row.get("base_agent_id"))
        and row.get("status") == "assigned"
    ]
    in_progress = [
        row for row in queue.get("tasks", [])
        if (str(row.get("id", "")).startswith("elastic-") or row.get("base_agent_id"))
        and row.get("status") == "in_progress"
    ]
    return {
        "runner": state,
        "canonical_runs": canonical,
        "executor": _runner_executor_status(),
        "background": _elastic_runner_status(),
        "queue": {
            "assigned": len(assigned),
            "in_progress": len(in_progress),
            "assigned_tasks": assigned,
            "in_progress_tasks": in_progress,
        },
        "source_of_truth": canonical["source"],
        "updated_at": canonical.get("updated_at") or state.get("updated_at") or _now_utc(),
    }


@router.get("/work-runs")
def list_work_runs(
    project_id: str = "",
    task_id: str = "",
    status: str = "",
    limit: int = Query(100, ge=1, le=500),
):
    try:
        rows = work_run_service.list(
            project_id=project_id,
            task_id=task_id,
            status=status,
            limit=limit,
        )
    except InvalidWorkTransition as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {
        "runs": rows,
        "total": len(rows),
        "source": "unified_dashboard.db:work_runs",
    }


@router.get("/work-runs/{run_id}")
def get_work_run(run_id: str):
    try:
        return work_run_service.get(run_id)
    except WorkRunNotFound as exc:
        raise HTTPException(status_code=404, detail="Work run not found") from exc


@router.post("/work-runs/{run_id}/transition")
def transition_work_run(run_id: str, req: WorkRunTransitionRequest):
    try:
        return work_run_service.transition(
            run_id,
            req.status,
            actor=req.actor,
            detail=req.detail,
            result_summary=req.result_summary,
        )
    except WorkRunNotFound as exc:
        raise HTTPException(status_code=404, detail="Work run not found") from exc
    except InvalidWorkTransition as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc


@router.post("/agents/runner/tick")
def tick_agent_runner(req: AgentRunnerTickRequest = AgentRunnerTickRequest()):
    return _runner_tick(req)


@router.post("/agents/instances/{instance_id}/claim")
def claim_agent_instance(instance_id: str, execute: bool = False):
    return _runner_tick(AgentRunnerTickRequest(instance_id=instance_id, execute=execute))


@router.post("/agents/spawn", status_code=201)
def spawn_agent_instance(req: AgentSpawnRequest):
    templates = _template_by_id()
    base_agent_id = _safe_id(req.base_agent_id)
    requested_by = _safe_id(req.requested_by)
    template = templates.get(base_agent_id)
    if not template:
        raise HTTPException(status_code=404, detail="Agent template not found")
    if not template.get("spawnable", False):
        raise HTTPException(status_code=400, detail="Agent template is not spawnable")

    instance_id = _safe_id(req.instance_id or f"{requested_by}-{base_agent_id}")
    store = _load_agent_instance_store()
    now = _now_utc()
    existing = next((row for row in store["instances"] if row.get("id") == instance_id), None)
    payload = {
        "id": instance_id,
        "instance_id": instance_id,
        "type": "elastic",
        "status": "running" if req.task_id else "idle",
        "base_agent_id": base_agent_id,
        "base_agent_name": template.get("name", base_agent_id),
        "requested_by": requested_by,
        "project_id": req.project_id,
        "task_id": req.task_id,
        "development_point_id": req.development_point_id,
        "role": template.get("role", ""),
        "capabilities": template.get("capabilities", []),
        "reason": req.reason,
        "release_policy": req.release_policy or template.get("default_release_policy", "task_completed"),
        "updated_at": now,
        "source": "agent-instance-registry",
    }
    if existing:
        existing.update(payload)
        existing.setdefault("created_at", now)
        instance = existing
        action = "updated"
    else:
        payload["created_at"] = now
        store["instances"].append(payload)
        instance = payload
        action = "created"
    dispatch = None
    if instance.get("task_id"):
        dispatch = _dispatch_instance_to_openclaw(instance, template)
        instance["dispatch_status"] = dispatch["action"]
        instance["dispatch_task_id"] = dispatch["queue_task"]["id"]
        instance["workspace"] = dispatch["workspace"]
    _save_agent_instance_store(store)
    return {"action": action, "instance": _enrich_agent_instance(instance), "dispatch": dispatch}


@router.post("/agents/instances/{instance_id}/release")
def release_agent_instance(instance_id: str, req: AgentReleaseRequest):
    store = _load_agent_instance_store()
    row = next((item for item in store["instances"] if item.get("id") == instance_id), None)
    if not row:
        raise HTTPException(status_code=404, detail="Agent instance not found")
    row["status"] = "released"
    row["released_by"] = req.released_by
    row["release_reason"] = req.reason
    row["released_at"] = _now_utc()
    row["updated_at"] = row["released_at"]
    queue_task = _mark_dispatch_instance(row, "released", req.reason)
    canonical_run = work_run_service.latest_for_dispatch(_instance_dispatch_id(row))
    if canonical_run and canonical_run.get("status") not in {"completed", "cancelled"}:
        try:
            canonical_run = work_run_service.transition(
                canonical_run["id"],
                "cancelled",
                actor=req.released_by,
                detail=req.reason,
            )
        except InvalidWorkTransition:
            pass
    _save_agent_instance_store(store)
    return {"instance": _enrich_agent_instance(row), "queue_task": queue_task, "work_run": canonical_run}


@router.post("/agents/instances/{instance_id}/complete")
def complete_agent_instance(instance_id: str, req: AgentCompleteRequest):
    store = _load_agent_instance_store()
    row = next((item for item in store["instances"] if item.get("id") == instance_id), None)
    if not row:
        raise HTTPException(status_code=404, detail="Agent instance not found")
    now = _now_utc()
    row["status"] = "completed"
    row["completed_by"] = req.completed_by
    row["completed_at"] = now
    row["updated_at"] = now
    row["result_summary"] = req.result_summary
    queue_task = _mark_dispatch_instance(row, "done", req.result_summary)
    canonical_run = work_run_service.latest_for_dispatch(_instance_dispatch_id(row))
    if canonical_run and canonical_run.get("status") in {"review", "verifying"}:
        canonical_run = work_run_service.transition(
            canonical_run["id"],
            "completed",
            actor=req.completed_by,
            detail="Legacy agent completion accepted as review approval",
            result_summary=req.result_summary,
        )
    commit = {
        "id": f"mem-{instance_id}-{len(store.get('memory_commits', [])) + 1}",
        "instance_id": instance_id,
        "base_agent_id": row.get("base_agent_id", ""),
        "project_id": row.get("project_id", ""),
        "task_id": row.get("task_id", ""),
        "development_point_id": row.get("development_point_id", ""),
        "summary": req.memory_summary or req.result_summary,
        "decisions": req.decisions,
        "blockers": req.blockers,
        "next_actions": req.next_actions,
        "committed_by": req.completed_by,
        "created_at": now,
        "target_memory": ["project", "base_agent"],
    }
    store.setdefault("memory_commits", []).append(commit)
    _save_agent_instance_store(store)
    return {
        "instance": _enrich_agent_instance(row),
        "memory_commit": commit,
        "queue_task": queue_task,
        "work_run": canonical_run,
    }


@router.post("/agents/release")
def release_agent_instance_alias(instance_id: str = Query(..., min_length=1), req: AgentReleaseRequest = AgentReleaseRequest()):
    return release_agent_instance(instance_id, req)


@router.get("/agents/status")
def list_agent_project_status():
    rows = _agent_rows()
    return {"agents": rows, "total": len(rows), "source": "projects-v3"}


@router.get("/agents/dashboard")
def list_agent_dashboard():
    agents, aggregate_source, observed_at = _load_agent_sources()
    skill_summary = skill_manager.skill_summary_by_agent()
    work_rows = _agent_rows(agents)
    work_by_agent = {row.get("agent_id"): row for row in work_rows}
    dashboard = []

    for agent in agents:
        aid = agent.get("id", "")
        work = work_by_agent.get(aid, {})
        last_seen = agent.get("last_seen") or agent.get("last_heartbeat") or agent.get("updated_at")
        stale = not bool(last_seen)
        stale_reason = "" if last_seen else "missing-last-seen"
        current_work = None
        if work.get("current_task_id"):
            current_work = {
                "project_id": work.get("current_project_id"),
                "project_name": work.get("current_project_name"),
                "task_id": work.get("current_task_id"),
                "task_title": work.get("current_task_title"),
                "development_point_id": work.get("current_development_point_id"),
                "development_point_title": work.get("current_development_point_title"),
                "task_progress": work.get("task_progress", 0),
                "project_progress": work.get("project_progress", 0),
                "source": "projects-v3",
            }

        status = work.get("status") or agent.get("status", "unknown")
        profile_source = agent.get("_dashboard_source", aggregate_source)
        agent_skill_summary = skill_summary.get(aid, {"total": 0, "skills": [], "issues": 0})
        dashboard.append({
            "id": aid,
            "agent_id": aid,
            "name": agent.get("name", aid),
            "agent_name": agent.get("name", aid),
            "role": agent.get("role", ""),
            "team": agent.get("team", ""),
            "memory": agent.get("memory", []),
            "status": status,
            "current_task": work.get("current_task_title") or agent.get("current_task", ""),
            "legacy_current_task": agent.get("current_task", ""),
            "current_project_id": work.get("current_project_id"),
            "current_project_name": work.get("current_project_name"),
            "current_task_id": work.get("current_task_id"),
            "current_task_title": work.get("current_task_title"),
            "current_development_point_id": work.get("current_development_point_id"),
            "current_development_point_title": work.get("current_development_point_title"),
            "task_progress": work.get("task_progress", 0),
            "project_progress": work.get("project_progress", 0),
            "updated_at": observed_at,
            "last_seen": last_seen,
            "stale": stale,
            "stale_reason": stale_reason,
            "profile": {
                "id": aid,
                "name": agent.get("name", aid),
                "role": agent.get("role", ""),
                "team": agent.get("team", ""),
                "skills": agent.get("skills", []),
                "memory": agent.get("memory", []),
            },
            "runtime": {
                "status": status,
                "raw_status": agent.get("status", "unknown"),
                "current_task": agent.get("current_task", ""),
                "last_seen": last_seen,
                "updated_at": observed_at,
                "stale": stale,
                "stale_reason": stale_reason,
            },
            "current_work": current_work,
            "skill_summary": agent_skill_summary,
            "skills": agent_skill_summary.get("skills", []),
            "source": {
                "profile": profile_source,
                "runtime": profile_source,
                "work": "projects-v3" if current_work else "none",
                "skills": "skill-registry",
            },
        })

    return {
        "agents": dashboard,
        "total": len(dashboard),
        "source": "agents-dashboard",
        "agent_source": aggregate_source,
        "updated_at": observed_at,
    }


@router.get("/agents/organization")
def get_agent_organization():
    try:
        document = unified_data_manager.get_agent_organization_document()
        if document.get("root"):
            return document
    except Exception:
        pass
    if not AGENT_ORGANIZATION_FILE.exists():
        raise HTTPException(status_code=404, detail="Agent organization data not found")
    try:
        document = json.loads(AGENT_ORGANIZATION_FILE.read_text(encoding="utf-8"))
        document["source"] = "agent-organization-json"
        return document
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=500, detail=f"Invalid agent organization data: {exc}") from exc


@router.get("/agents/{agent_id}/current-work")
def get_agent_current_work(agent_id: str):
    for row in _agent_rows():
        if row.get("agent_id") == agent_id:
            return row
    raise HTTPException(status_code=404, detail="Agent not found")


@router.get("/agents/{agent_id}/work-items")
def get_agent_work_items(agent_id: str):
    return project_manager.get_agent_work_items(agent_id)
